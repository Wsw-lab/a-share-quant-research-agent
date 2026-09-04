from pathlib import Path
import unittest

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = (
    ROOT
    / "docs/working-paper/A_Share_Factor_Specification_Effects_Stage1_Manuscript.docx"
)
RESULT_TABLES = (
    ROOT
    / "studies/pit_factor_bias_decomposition_v2/prespecified_results_tables.md"
)


class Stage1ManuscriptArtifactTest(unittest.TestCase):
    def test_prespecified_results_supplement_has_complete_rows(self) -> None:
        text = RESULT_TABLES.read_text(encoding="utf-8")

        self.assertNotIn(r"\n|", text)
        self.assertEqual(text.count("\n| PRIMARY |"), 1)
        self.assertEqual(text.count("\n| SECONDARY |"), 28)
        self.assertEqual(text.count("\n| CELL |"), 72)
        self.assertEqual(text.count("`C_publication_isolation_momentum_60d`"), 1)
        self.assertEqual(text.count("`C_publication_isolation_low_vol_20d`"), 1)

    def test_docx_contains_result_shell_and_uses_plain_black_heading_system(self) -> None:
        """Catch omission of the fixed result shell or return of decorative headings."""

        doc = Document(DOCX)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn("Pre-specified historical result shell", text)
        self.assertIn("Not yet estimated", text)

        title = next(
            paragraph for paragraph in doc.paragraphs
            if paragraph.style.name == "Title"
        )
        title_borders = (
            title._p.pPr.find(qn("w:pBdr"))
            if title._p.pPr is not None
            else None
        )
        self.assertIsNone(title_borders)

        for style_name in (
            "Title",
            "Subtitle",
            "Heading 1",
            "Heading 2",
            "Heading 3",
        ):
            color = doc.styles[style_name].font.color.rgb
            self.assertEqual(str(color), "000000")


if __name__ == "__main__":
    unittest.main()
