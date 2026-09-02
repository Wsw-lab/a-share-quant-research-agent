#!/usr/bin/env python3
"""Build the anonymous Stage-1 working paper from its Markdown source.

The builder deliberately starts from the retained protocol DOCX so that page
geometry, theme, styles, and header/footer conventions remain consistent with
the existing working-paper package.  It does not run any empirical analysis.
"""

from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = Path(__file__).with_name("stage1_manuscript.md")
TEMPLATE = ROOT / "docs/working-paper/A_Share_Stage2_Registered_Report_Protocol.docx"
OUTPUT = ROOT / "docs/working-paper/A_Share_Factor_Specification_Effects_Stage1_Manuscript.docx"
ASSET_DIR = ROOT / "docs/working-paper/generated-assets"

NAVY = "0B2545"
BLUE = "2E74B5"
MID_BLUE = "4F81BD"
PALE_BLUE = "E9EFF5"
PALE_GOLD = "FFF4CC"
GOLD = "B98500"
GREEN = "2F6B4F"
PALE_GREEN = "E7F2EC"
GREY = "5A6573"
LIGHT_GREY = "F3F5F7"
BORDER = "B8C5D3"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width_inches: float) -> None:
    width = Inches(width_inches)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_grid(table, widths: list[float]) -> None:
    """Bind a fixed grid because cell widths alone are not honored everywhere."""
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(sum(int(Inches(width).twips) for width in widths)))
    tbl_width.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(Inches(width).twips)))
        grid.append(col)


def set_bottom_rule(paragraph, color=BLUE, size="14") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_field(run, instruction: str, placeholder: str) -> None:
    run._r.clear_content()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instruction} "
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    for child in (fld_char_begin, instr_text, fld_char_sep, text, fld_char_end):
        run._r.append(child)


def new_numbering_id(doc: Document) -> int:
    """Create a new real numbered-list instance that restarts at one."""
    numbering = doc.part.numbering_part.element
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), "7")
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = rgb("202830")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.04
    normal.paragraph_format.space_after = Pt(3.2)
    normal.paragraph_format.widow_control = True

    for style_name, size, before, after in (
        ("Heading 1", 13.6, 7, 3.5),
        ("Heading 2", 11.2, 5, 2.5),
        ("Heading 3", 10.2, 4, 2),
    ):
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(BLUE)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.keep_together = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(14.2)
    subtitle.font.color.rgb = rgb("1F4D78")
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(8.6)
    caption.font.bold = True
    caption.font.color.rgb = rgb(MID_BLUE)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(2)
    caption.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(9.6)
        st.paragraph_format.left_indent = Inches(0.22)
        st.paragraph_format.first_line_indent = Inches(-0.16)
        st.paragraph_format.space_after = Pt(1.6)
        st.paragraph_format.line_spacing = 1.02

    if "Protocol Kicker" in styles:
        kicker = styles["Protocol Kicker"]
        kicker.font.name = "Calibri"
        kicker.font.size = Pt(9.5)
        kicker.font.bold = True
        kicker.font.color.rgb = rgb(MID_BLUE)
        kicker.paragraph_format.space_after = Pt(5)


def configure_section_and_headers(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.30)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("PRE-RESULTS PROTOCOL   |   A-SHARE ROE INFORMATION TIMING")
    r.font.name = "Calibri"
    r.font.size = Pt(7.5)
    r.font.bold = True
    r.font.color.rgb = rgb(GREY)
    set_bottom_rule(p, color=BORDER, size="4")

    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lead = p.add_run("ANONYMOUS DRAFT · NOT REGISTERED   |   SEPTEMBER 2026   |   ")
    lead.font.name = "Calibri"
    lead.font.size = Pt(7.3)
    lead.font.color.rgb = rgb(GREY)
    page = p.add_run()
    set_field(page, "PAGE", "1")
    page.font.name = "Calibri"
    page.font.size = Pt(7.3)
    page.font.color.rgb = rgb(GREY)
    mid = p.add_run(" / ")
    mid.font.size = Pt(7.3)
    mid.font.color.rgb = rgb(GREY)
    pages = p.add_run()
    set_field(pages, "NUMPAGES", "1")
    pages.font.name = "Calibri"
    pages.font.size = Pt(7.3)
    pages.font.color.rgb = rgb(GREY)

    first_header = section.first_page_header
    fp = first_header.paragraphs[0] if first_header.paragraphs else first_header.add_paragraph()
    fp.clear()
    first_footer = section.first_page_footer
    fp = first_footer.paragraphs[0] if first_footer.paragraphs else first_footer.add_paragraph()
    fp.clear()


INLINE_RE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+`)")


def add_inline(paragraph, text: str, *, base_size: float | None = None, base_color: str | None = None) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            if base_size:
                run.font.size = Pt(base_size)
            if base_color:
                run.font.color.rgb = rgb(base_color)
        token = match.group(0)
        if token.startswith("**"):
            content = token[2:-2]
            run = paragraph.add_run(content)
            run.bold = True
        elif token.startswith("*"):
            content = token[1:-1]
            run = paragraph.add_run(content)
            run.italic = True
        else:
            content = token[1:-1]
            run = paragraph.add_run(content)
            run.font.name = "Courier New"
            run.font.size = Pt((base_size or 9.6) - 0.4)
            run.font.color.rgb = rgb(NAVY)
        if base_size and not token.startswith("`"):
            run.font.size = Pt(base_size)
        if base_color:
            run.font.color.rgb = rgb(base_color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        if base_size:
            run.font.size = Pt(base_size)
        if base_color:
            run.font.color.rgb = rgb(base_color)


def add_cover(doc: Document, meta: dict[str, str]) -> None:
    p = doc.add_paragraph(style="Protocol Kicker" if "Protocol Kicker" in doc.styles else "Normal")
    p.add_run("ANONYMOUS PRE-RESULTS PROTOCOL DRAFT")

    p = doc.add_paragraph(style="Title")
    p.add_run(meta["title"])
    set_bottom_rule(p, color=BLUE, size="18")

    p = doc.add_paragraph(style="Subtitle")
    p.add_run(meta["subtitle"])

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Anonymous manuscript for editorial evaluation")
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = rgb(GREY)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    set_table_borders(table, color="E0B84E", size="6")
    cell = table.cell(0, 0)
    set_cell_width(cell, 7.0)
    set_cell_shading(cell, PALE_GOLD)
    set_cell_margins(cell, top=130, start=170, bottom=130, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("EVIDENCE STATUS")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = rgb(GOLD)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_inline(
        p,
        "The 2025–2026 pilot is observed; its aggregate results and known limitations are publicly disclosed. The 2010–2022 historical confirmation is blocked for data feasibility, has not been externally registered or authorized, and has not been executed.",
        base_size=9.6,
        base_color=NAVY,
    )

    doc.add_paragraph()
    details = [
        ("Date", meta["date"]),
        ("Keywords", meta["keywords"]),
        ("JEL classifications", meta["jel"]),
        ("Repository state", "Public verification package; licensed empirical rows are not redistributed"),
    ]
    table = doc.add_table(rows=len(details), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    set_table_borders(table, color=BORDER, size="4")
    for idx, (label, value) in enumerate(details):
        row = table.rows[idx]
        set_cell_width(row.cells[0], 1.35)
        set_cell_width(row.cells[1], 5.65)
        for cell in row.cells:
            set_cell_margins(cell, top=75, start=110, bottom=75, end=110)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], PALE_BLUE)
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label.upper())
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = rgb(BLUE)
        p = row.cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        add_inline(p, value, base_size=8.8, base_color="27313A")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("CLAIM BOUNDARY")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = rgb(BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_inline(
        p,
        "No confirmatory result, performance claim, trading-readiness claim, historical-vintage claim, or claim of journal acceptance appears in this manuscript.",
        base_size=9.2,
        base_color=GREY,
    )

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def font_path(bold: bool = False) -> str:
    candidates = (
        ["/System/Library/Fonts/Supplemental/Arial Bold.ttf", "/Library/Fonts/Arial Bold.ttf"]
        if bold
        else ["/System/Library/Fonts/Supplemental/Arial.ttf", "/Library/Fonts/Arial.ttf"]
    )
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return "/System/Library/Fonts/Supplemental/Arial.ttf"


def rounded_box(draw, box, fill, outline, radius=20, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def centered(draw, box, text, font, fill, spacing=4):
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    metrics = [draw.textbbox((0, 0), line, font=font) for line in lines]
    heights = [m[3] - m[1] for m in metrics]
    total = sum(heights) + spacing * (len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    for line, metric, height in zip(lines, metrics, heights):
        width = metric[2] - metric[0]
        draw.text((x1 + (x2 - x1 - width) / 2, y), line, font=font, fill=fill)
        y += height + spacing


def create_accounting_timeline_figure(path: Path) -> None:
    img = Image.new("RGB", (1800, 610), f"#{WHITE}")
    draw = ImageDraw.Draw(img)
    title = ImageFont.truetype(font_path(True), 38)
    head = ImageFont.truetype(font_path(True), 25)
    body = ImageFont.truetype(font_path(False), 22)
    small = ImageFont.truetype(font_path(False), 19)
    draw.text((65, 38), "Two clocks assign the same accounting record to different dates", font=title, fill=f"#{NAVY}")
    draw.line((65, 96, 1735, 96), fill=f"#{BLUE}", width=5)

    period_x, signal_x, publication_x, eligible_x = 210, 655, 1160, 1580
    timeline_y = 330
    draw.line((period_x, timeline_y, eligible_x, timeline_y), fill=f"#{BORDER}", width=7)
    draw.polygon([(eligible_x - 10, timeline_y - 15), (eligible_x + 18, timeline_y), (eligible_x - 10, timeline_y + 15)], fill=f"#{BORDER}")

    rounded_box(draw, (period_x, 145, publication_x, 245), f"#{PALE_GOLD}", f"#{GOLD}", radius=18, width=3)
    centered(
        draw,
        (period_x + 15, 155, publication_x - 15, 235),
        "PREMATURE-USE INTERVAL UNDER THE REPORT-PERIOD CLOCK",
        head,
        f"#{GOLD}",
    )

    milestones = [
        (period_x, "FISCAL PERIOD END", "Report-period rule\nfirst admits record", GOLD),
        (signal_x, "MONTHLY SIGNAL", "Record may be selected\nbefore publication", GOLD),
        (publication_x, "RECORDED PUBLICATION DATE", "Date stored in the\nsingle-version export", BLUE),
        (eligible_x, "FIRST CONSERVATIVE SESSION", "Publication rule admits record\nonly strictly after date", GREEN),
    ]
    for x, label, desc, color in milestones:
        draw.line((x, timeline_y - 18, x, timeline_y + 18), fill=f"#{color}", width=6)
        draw.ellipse((x - 10, timeline_y - 10, x + 10, timeline_y + 10), fill=f"#{color}")
        centered(draw, (x - 190, 355, x + 190, 402), label, small, f"#{color}")
        centered(draw, (x - 190, 410, x + 190, 490), desc, body, f"#{NAVY}", spacing=6)

    centered(
        draw,
        (95, 520, 1705, 585),
        "Database eligibility only · not a claim about the value investors first saw or every earlier public signal",
        body,
        f"#{GREY}",
    )
    img.save(path, dpi=(300, 300))


def create_design_figure(path: Path) -> None:
    img = Image.new("RGB", (1800, 720), f"#{WHITE}")
    draw = ImageDraw.Draw(img)
    title = ImageFont.truetype(font_path(True), 38)
    head = ImageFont.truetype(font_path(True), 27)
    body = ImageFont.truetype(font_path(False), 22)
    small = ImageFont.truetype(font_path(True), 20)
    draw.text((65, 35), "Pre-specified information and implementation design", font=title, fill=f"#{NAVY}")
    draw.line((65, 92, 1735, 92), fill=f"#{BLUE}", width=5)

    chain = [
        (80, 160, 450, 330, "A0", "Final survivors\nReport-period clock"),
        (530, 160, 900, 330, "A1", "Historical membership\nReport-period clock"),
        (980, 160, 1350, 330, "I0000", "Historical membership\nPublication-date clock"),
    ]
    for x1, y1, x2, y2, label, desc in chain:
        rounded_box(draw, (x1, y1, x2, y2), f"#{PALE_BLUE}", f"#{BLUE}", radius=22, width=4)
        centered(draw, (x1 + 15, y1 + 20, x2 - 15, y1 + 70), label, head, f"#{BLUE}")
        centered(draw, (x1 + 15, y1 + 78, x2 - 15, y2 - 20), desc, body, f"#{NAVY}", spacing=7)
    for x in (450, 900):
        draw.line((x, 245, x + 80, 245), fill=f"#{BORDER}", width=5)
        draw.polygon([(x + 67, 231), (x + 92, 245), (x + 67, 259)], fill=f"#{BORDER}")
    draw.line((1350, 245, 1415, 245), fill=f"#{BORDER}", width=5)
    draw.polygon([(1402, 231), (1427, 245), (1402, 259)], fill=f"#{BORDER}")
    centered(draw, (430, 115, 560, 155), "membership", small, f"#{GREY}")
    centered(draw, (875, 115, 1015, 155), "publication", small, f"#{GREY}")

    rounded_box(draw, (1415, 130, 1720, 355), f"#{PALE_GOLD}", f"#{GOLD}", radius=24, width=4)
    centered(draw, (1430, 150, 1705, 200), "PRIMARY", head, f"#{GOLD}")
    centered(draw, (1430, 205, 1705, 320), "ROE monthly IC\nI0000 − A1\n(two-sided inference)", body, f"#{NAVY}", spacing=7)

    draw.line((1165, 330, 1165, 410), fill=f"#{BLUE}", width=5)
    draw.polygon([(1151, 397), (1165, 422), (1179, 397)], fill=f"#{BLUE}")
    rounded_box(draw, (175, 430, 1625, 645), f"#{LIGHT_GREY}", f"#{BORDER}", radius=25, width=4)
    centered(draw, (210, 448, 1590, 490), "COMPLETE 16-VARIANT IMPLEMENTATION LATTICE", head, f"#{NAVY}")
    components = [
        (250, "ST\nexclusion"),
        (570, "Suspension\nexclusion"),
        (890, "CNY 5m\namount floor"),
        (1210, "One-session\nreturn lag"),
    ]
    for x, text in components:
        rounded_box(draw, (x, 515, x + 250, 610), f"#{WHITE}", f"#{MID_BLUE}", radius=16, width=3)
        centered(draw, (x + 10, 520, x + 240, 605), text, body, f"#{BLUE}", spacing=5)
    centered(draw, (180, 652, 1620, 704), "16 implementation variants × 4 factors = 64 implementation cells · exact monthly Shapley allocation", body, f"#{GREY}")
    img.save(path, dpi=(300, 300))


def parse_frontmatter(lines: list[str]) -> tuple[dict[str, str], list[str]]:
    if not lines or lines[0].strip() != "---":
        raise ValueError("Markdown front matter is required")
    end = lines.index("---", 1)
    meta: dict[str, str] = {}
    for line in lines[1:end]:
        key, value = line.split(":", 1)
        meta[key.strip()] = value.strip().strip('"')
    return meta, lines[end + 1 :]


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        values = [value.strip() for value in lines[index].strip().strip("|").split("|")]
        rows.append(values)
        index += 1
    if len(rows) < 2 or not all(re.fullmatch(r":?-{3,}:?", value) for value in rows[1]):
        raise ValueError(f"Malformed Markdown table at source line {start + 1}")
    return [rows[0], *rows[2:]], index


def table_widths(headers: list[str]) -> list[float]:
    key = tuple(headers)
    maps = {
        ("Variant", "Factor", "Mean IC", "NW t", "Top minus universe", "Mean N"): [1.08, 1.45, 0.78, 0.70, 1.73, 1.26],
        ("Factor", "Historical-universe report-end IC", "Recorded-publication IC", "Difference", "Role"): [1.12, 1.88, 1.52, 0.82, 1.66],
        ("Group", "Count", "Definition", "Inference", "Authorized role"): [1.05, 0.52, 2.10, 1.55, 1.78],
        ("Observed condition", "Authorized interpretation", "Prohibited interpretation"): [1.55, 2.75, 2.70],
        ("ID", "Historical universe", "Accounting availability", "Enabled implementation components"): [2.35, 1.25, 1.35, 2.05],
    }
    if key in maps:
        return maps[key]
    return [7.0 / len(headers)] * len(headers)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    headers = rows[0]
    widths = table_widths(headers)
    table = doc.add_table(rows=len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_grid(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    numeric_headers = {
        "Count",
        "Mean IC",
        "NW t",
        "Top minus universe",
        "Mean N",
        "Historical-universe report-end IC",
        "Recorded-publication IC",
        "Difference",
    }
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        keep_row_together(row)
        for col_index, value in enumerate(values):
            cell = row.cells[col_index]
            set_cell_width(cell, widths[col_index])
            compact = headers[0] == "Group"
            set_cell_margins(
                cell,
                top=38 if compact else 58,
                start=62 if compact else 72,
                bottom=38 if compact else 58,
                end=62 if compact else 72,
            )
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index == 0:
                set_cell_shading(cell, PALE_BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            if headers[col_index] in numeric_headers:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            font_size = 7.15 if compact else (7.55 if len(headers) >= 5 else 7.8)
            add_inline(paragraph, value, base_size=font_size, base_color="26313B")
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = rgb(NAVY)

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    after.paragraph_format.space_before = Pt(0)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    set_table_borders(table, color=GOLD, size="5")
    cell = table.cell(0, 0)
    set_cell_width(cell, 7.0)
    set_cell_shading(cell, PALE_GOLD)
    set_cell_margins(cell, top=100, start=130, bottom=100, end=130)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    add_inline(paragraph, text, base_size=9.0, base_color=NAVY)
    for run in paragraph.runs[:1]:
        run.bold = True


def add_figure(doc: Document, path: Path, width: float, alt_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", alt_text)
    picture._inline.docPr.set("title", alt_text)


def add_body(doc: Document, lines: list[str], figures: dict[str, Path]) -> None:
    in_references = False
    active_numbering_id: int | None = None
    index = 0
    while index < len(lines):
        raw = lines[index].rstrip()
        text = raw.strip()
        if not text:
            active_numbering_id = None
            index += 1
            continue

        if text.startswith("[[FIGURE:"):
            key = text.removeprefix("[[FIGURE:").removesuffix("]]" )
            alt_texts = {
                "ACCOUNTING_TIMELINE": "Timeline contrasting fiscal-period eligibility with conservative eligibility strictly after the recorded publication date.",
                "DESIGN_MAP": "Pre-specified information-set chain and complete 16-variant implementation lattice, producing 64 implementation factor-variant cells.",
            }
            add_figure(doc, figures[key], 6.85, alt_texts[key])
            index += 1
            continue

        if text == "[[PAGEBREAK]]":
            doc.add_page_break()
            index += 1
            continue

        if text.startswith("|") and index + 1 < len(lines) and lines[index + 1].strip().startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue

        if text.startswith("# "):
            heading = text[2:]
            if heading == "Abstract":
                p = doc.add_paragraph(style="Heading 1")
                p.add_run(heading)
            else:
                if heading == "References":
                    p_break = doc.add_paragraph()
                    p_break.add_run().add_break(WD_BREAK.PAGE)
                    in_references = True
                p = doc.add_paragraph(style="Heading 1")
                p.add_run(heading)
            index += 1
            continue

        if text.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            p.add_run(text[3:])
            index += 1
            continue

        if text.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            p.add_run(text[4:])
            index += 1
            continue

        if text.startswith("> "):
            add_callout(doc, text[2:])
            index += 1
            continue

        if text.startswith("**Table ") and text.endswith("**"):
            p = doc.add_paragraph(style="Caption")
            add_inline(p, text)
            index += 1
            continue

        if text.startswith("*Notes.*") or text.startswith("*Figure "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_together = True
            add_inline(p, text, base_size=7.9, base_color=GREY)
            index += 1
            continue

        if re.match(r"^\d+\.\s", text):
            if active_numbering_id is None:
                active_numbering_id = new_numbering_id(doc)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, active_numbering_id)
            add_inline(p, re.sub(r"^\d+\.\s+", "", text))
            index += 1
            continue

        if text.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, text[2:])
            index += 1
            continue

        p = doc.add_paragraph()
        if in_references:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Inches(0.20)
            p.paragraph_format.first_line_indent = Inches(-0.20)
            p.paragraph_format.space_after = Pt(2.2)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, text, base_size=8.25, base_color="29323B")
        elif text.startswith("`") and text.endswith("`") and text.count("`") == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, text, base_size=9.0, base_color=NAVY)
        else:
            add_inline(p, text)
        index += 1


def set_document_properties(doc: Document, meta: dict[str, str]) -> None:
    props = doc.core_properties
    props.title = meta["title"]
    props.subject = meta["subtitle"]
    props.author = ""
    props.last_modified_by = ""
    props.keywords = meta["keywords"]
    props.comments = "Anonymous pre-results Stage-1 working paper; confirmatory outcomes not accessed."
    props.category = "Working paper"


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    accounting_timeline = ASSET_DIR / "accounting_availability_timeline.png"
    design_figure = ASSET_DIR / "prespecified_design_map.png"
    create_accounting_timeline_figure(accounting_timeline)
    create_design_figure(design_figure)

    meta, body = parse_frontmatter(SOURCE.read_text(encoding="utf-8").splitlines())
    doc = Document(TEMPLATE)
    clear_document_body(doc)
    configure_styles(doc)
    configure_section_and_headers(doc)
    set_document_properties(doc, meta)
    add_cover(doc, meta)
    add_body(doc, body, {"ACCOUNTING_TIMELINE": accounting_timeline, "DESIGN_MAP": design_figure})

    # Encourage Word/LibreOffice to refresh PAGE and NUMPAGES fields when opened.
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
