"""Build the three provider-outreach DOCX artifacts.

The two letters use the standard_business_brief preset. The English project
brief uses the grant_proposal alias of narrative_proposal. Identity fields are
deliberate placeholders because the repository does not establish the owner's
current student or institutional status.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "docs" / "provider-outreach"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
GRAY = RGBColor(89, 96, 105)
BLACK = RGBColor(0, 0, 0)
PLACEHOLDER = "FFF2CC"


def set_run_font(
    run,
    *,
    size: float = 11,
    color: RGBColor = BLACK,
    bold: bool = False,
    italic: bool = False,
    east_asia: str = "PingFang SC",
) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def highlight_placeholder(run) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), PLACEHOLDER)
    run._element.get_or_add_rPr().append(shading)


def set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    """Set fixed DXA widths; widths must sum to the 9360-DXA content width."""
    if sum(widths) != 9360:
        raise ValueError("table widths must sum to 9360 DXA")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def configure_document(doc: Document, *, proposal: bool = False) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25 if proposal else 1.10
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    doc.core_properties.author = "A-Share Research Project"
    doc.core_properties.last_modified_by = "A-Share Research Project"
    doc.core_properties.comments = "Provider outreach materials; identity placeholders require completion before dispatch."


def add_header_footer(doc: Document, left: str, right: str) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{left}  |  {right}")
    set_run_font(run, size=9, color=GRAY, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fr = fp.add_run("CONFIDENTIAL DRAFT - Complete highlighted identity fields before sending")
    set_run_font(fr, size=8.5, color=GRAY, italic=True)


def add_title(doc: Document, kicker: str, title: str, subtitle: str) -> None:
    kp = doc.add_paragraph()
    kp.paragraph_format.space_after = Pt(4)
    kr = kp.add_run(kicker.upper())
    set_run_font(kr, size=9.5, color=BLUE, bold=True)

    tp = doc.add_paragraph()
    tp.paragraph_format.space_after = Pt(4)
    tp.paragraph_format.keep_with_next = True
    tr = tp.add_run(title)
    set_run_font(tr, size=22, color=INK, bold=True)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(12)
    sp.paragraph_format.keep_with_next = True
    sr = sp.add_run(subtitle)
    set_run_font(sr, size=12, color=GRAY, italic=True)


def add_metadata(doc: Document, provider: str) -> None:
    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, [4680, 4680])
    set_repeat_table_header(table.rows[0])
    values = [
        ("To", f"{provider} Academic / Education Partnerships Team"),
        ("From", "[Your name - complete before sending]"),
        ("Status", "[Student / early-career researcher / independent researcher - choose truthfully]"),
        ("Affiliation", "[Institution or Independent Researcher]"),
        ("Email", "[Reply email]"),
        ("Date", "[Date sent]"),
    ]
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            label, value = values[row_index * 2 + col_index]
            set_cell_fill(cell, "F2F4F7")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            label_run = p.add_run(f"{label}: ")
            set_run_font(label_run, size=9.5, color=DARK_BLUE, bold=True)
            value_run = p.add_run(value)
            set_run_font(value_run, size=9.5, color=BLACK)
            if value.startswith("["):
                highlight_placeholder(value_run)
def add_labeled_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = True
    lr = p.add_run(f"{label} ")
    set_run_font(lr, bold=True, color=DARK_BLUE)
    rr = p.add_run(text)
    set_run_font(rr)


def build_provider_letter(provider: str, provider_specific: str, filename: str) -> Path:
    doc = Document()
    configure_document(doc)
    add_header_footer(doc, "Academic Data Partnership", provider)
    add_title(
        doc,
        "Outcome-blind research collaboration",
        "合作申请：预注册 A 股实证研究的数据访问资助",
        "In-kind data sponsorship and technical collaboration request",
    )
    add_metadata(doc, provider)

    p = doc.add_paragraph()
    r = p.add_run(f"尊敬的 {provider} 学术合作团队：")
    set_run_font(r, bold=True)

    body = (
        "我们正在准备英文论文《Report Dates, Publication Dates, and the A-Share ROE Signal: "
        "A Pre-Specified Historical Confirmation》。研究问题是：信息可得时间与实施惯例会在多大程度上扭曲 A 股因子证据？"
        "公开代码与试点已经存在，但 2010-2022 年 Stage-2 尚未运行，历史因子收益、IC 和策略排序仍保持封存。"
    )
    p = doc.add_paragraph(body)
    p.paragraph_format.keep_together = True
    for run in p.runs:
        set_run_font(run)

    add_labeled_paragraph(
        doc,
        "合作请求。",
        "我们希望申请非现金、限期的数据访问资助及一名技术联络人，用于核实 2009 年 warm-up、2010-2022 年样本和 2023 年 1 月收益端点所需的行情、复权、上市/退市、实际财报披露日期、ST、停牌、成交额与交易日历字段。第一轮只需要字段字典、覆盖说明、报价/资助条件和书面授权答复，不需要发送任何结果样本或原始数据。",
    )
    add_labeled_paragraph(doc, f"为何联系 {provider}。", provider_specific)
    add_labeled_paragraph(
        doc,
        "合作价值。",
        "若数据与权利审查通过，我们将在论文中规范引用并披露数据支持，在不公开授权原始数据的前提下开放研究设计、代码、合成测试、哈希与获准的汇总覆盖信息，并可在论文完成后分享方法与非敏感的数据质量反馈。该合作可形成一个强调预注册、全结果披露和可复核性的独立学术应用案例。",
    )
    add_labeled_paragraph(
        doc,
        "研究独立性。",
        "资助不以正面或显著结果为条件；供应商不拥有分析、结论、投稿或发表的否决权，只可在投稿前核对产品名称、字段定义和数据来源表述。任何资助、数据支持或潜在利益冲突都会如实披露；仅提供数据访问或一般技术支持不自动产生共同作者资格。",
    )
    add_labeled_paragraph(
        doc,
        "必须确认的权利。",
        "书面许可需要覆盖加密本地研究、汇总统计发表、供应商及产品名称公开、字段映射引用、文件哈希与覆盖率披露，以及通过供应商授权或持牌环境开展受控审稿复核。若这些条件无法满足，我们不会将该数据用于 Stage-2。",
    )

    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(4)
    closing.paragraph_format.keep_together = True
    for text, bold in (
        ("项目仓库：", True),
        ("https://github.com/Wsw-lab/a-share-quant-research-agent\n", False),
        ("随信材料：", True),
        ("一页英文项目简介、字段能力表、权利确认表与私有交付说明。\n", False),
        ("此致  ", False),
        ("[Your name] | [Affiliation / Independent Researcher] | [Reply email]", False),
    ):
        run = closing.add_run(text)
        set_run_font(run, bold=bold)
        if text.startswith("["):
            highlight_placeholder(run)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


def add_brief_section(doc: Document, heading: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(heading.upper())
    set_run_font(r, size=10, color=BLUE, bold=True)
    body = doc.add_paragraph(text)
    body.paragraph_format.space_after = Pt(5)
    body.paragraph_format.keep_together = True
    for run in body.runs:
        set_run_font(run, size=10.5)


def build_project_brief() -> Path:
    doc = Document()
    configure_document(doc, proposal=True)
    add_header_footer(doc, "Research Partnership Brief", "A-Share Factor Evidence")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("IN-KIND DATA SPONSORSHIP PROPOSAL")
    set_run_font(r, size=9.5, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Report Dates, Publication Dates, and the A-Share ROE Signal")
    set_run_font(r, size=21, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("A pre-specified historical confirmation of information-timing and implementation effects")
    set_run_font(r, size=11.5, color=GRAY, italic=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(7)
    for text, placeholder in (
        ("Researcher: [Name]", True),
        ("  |  Status: [Current status]", True),
        ("  |  Affiliation: [Institution or Independent Researcher]", True),
        ("  |  Contact: [Email]", True),
    ):
        run = meta.add_run(text)
        set_run_font(run, size=9.2, color=BLACK)
        if placeholder:
            highlight_placeholder(run)

    add_brief_section(
        doc,
        "Research question",
        "How much do information timing and implementation conventions distort A-share factor evidence? The study separates the effect of using fiscal report-period dates from recorded publication dates and then decomposes four implementation choices: ST exclusion, suspension exclusion, a 20-session liquidity floor, and a one-session lag.",
    )
    add_brief_section(
        doc,
        "Why it matters",
        "A factor can appear predictive when accounting information is treated as available before it was published, or when historical membership, suspensions, liquidity, and execution timing are handled inconsistently. The project converts these often-hidden choices into a fully disclosed research object rather than presenting only the best-performing specification.",
    )
    add_brief_section(
        doc,
        "Pre-specified design",
        "The public 2025-June 2026 pilot is disclosed as prior evidence and will not be relabeled as confirmation. Stage 2 targets January 2010-December 2022, with a 2009 warm-up and exact forward endpoints through January 2023. It reports all 18 variants and 72 factor-variant cells. Confirmatory inference is limited to one primary estimand and a fixed 28-member secondary family with multiplicity control. Coverage is tested without viewing factor outcomes; the design is frozen and externally registered before a separately authorized execution.",
    )
    add_brief_section(
        doc,
        "Support requested",
        "Time-limited academic access to the required historical A-share datasets; exact product/table/field mapping and a technical liaison; and written permission for local non-commercial analysis, aggregate publication, provider and field citation, hashes, and controlled reviewer reruns. No result-bearing rows are requested during the initial review. The study remains BLOCKED_FOR_STAGE2 until coverage, rights, registration, and execution authorization all pass.",
    )
    add_brief_section(
        doc,
        "Value to the data partner",
        "A transparent academic use case; formal citation and acknowledgement; an open-source reproducibility framework that excludes licensed raw rows; permitted aggregate coverage and data-quality feedback; and an optional post-publication methods discussion. Support is not exchanged for positive findings, endorsement, authorship, or publication control. A single-version export will not be presented as revision or vintage history.",
    )

    repo = doc.add_paragraph()
    repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    repo.paragraph_format.space_before = Pt(5)
    repo.paragraph_format.space_after = Pt(0)
    r = repo.add_run("Public repository: https://github.com/Wsw-lab/a-share-quant-research-agent")
    set_run_font(r, size=9.5, color=DARK_BLUE, bold=True)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / "A_Share_Study_One_Page_Partner_Brief.docx"
    doc.save(path)
    return path


def main() -> None:
    paths = [
        build_provider_letter(
            "CSMAR",
            "CSMAR 的公开资料明确面向实证研究，并展示了股票交易、财务报表、特别处理及停复牌等研究数据方向；贵方也有支持青年学者、论文活动和数据库使用权的历史实践。因此，我们希望首先探讨以限期数据权限和技术支持为主的学术合作，而不是以研究结论为交换的商业宣传。",
            "CSMAR_Academic_Data_Sponsorship_Request.docx",
        ),
        build_provider_letter(
            "RESSET",
            "RESSET 的公开资料以实证研究、高校教学和量化研究为主要应用方向，并明确要求研究论文规范引用数据来源。我们的预注册、全规格披露和授权审稿复核安排与这一学研定位相匹配，因此希望探讨限期研究权限、技术联络和书面发表/复核授权。",
            "RESSET_Academic_Data_Sponsorship_Request.docx",
        ),
        build_project_brief(),
    ]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
